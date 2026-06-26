from __future__ import annotations

import csv
import json
from collections.abc import Iterable
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

import yaml
from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from pypdf import PdfReader

from ..utils import normalise_text, sha256_text, stable_hash

PARSER_VERSION = "filesystem-v2"
SUPPORTED_SUFFIXES = {
    ".txt",
    ".md",
    ".markdown",
    ".json",
    ".jsonl",
    ".csv",
    ".tsv",
    ".log",
    ".yaml",
    ".yml",
    ".html",
    ".htm",
    ".pdf",
    ".docx",
}


class UnsupportedDocumentError(ValueError):
    pass


@dataclass(frozen=True)
class ParsedDocument:
    source: str
    source_uri: str
    title: str
    text: str
    content_hash: str
    ingest_key: str
    parser_version: str = PARSER_VERSION
    metadata: dict[str, Any] = field(default_factory=dict)


def iter_supported_files(path: Path) -> Iterable[Path]:
    candidates = [path] if path.is_file() else path.rglob("*")
    for candidate in sorted(candidates):
        if candidate.is_file() and candidate.suffix.casefold() in SUPPORTED_SUFFIXES:
            yield candidate


def file_fingerprint(path: Path) -> str:
    stat = path.stat()
    return stable_hash(path.resolve().as_posix(), stat.st_size, stat.st_mtime_ns)


def _read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


def _read_json(path: Path) -> str:
    raw = _read_text(path)
    if path.suffix.casefold() == ".jsonl":
        rows: list[Any] = []
        for line in raw.splitlines():
            if line.strip():
                rows.append(json.loads(line))
        return json.dumps(rows, ensure_ascii=False, indent=2)
    return json.dumps(json.loads(raw), ensure_ascii=False, indent=2)


def _read_yaml(path: Path) -> str:
    return json.dumps(yaml.safe_load(_read_text(path)), ensure_ascii=False, indent=2)


def _read_delimited(path: Path) -> str:
    delimiter = "\t" if path.suffix.casefold() == ".tsv" else ","
    with path.open(encoding="utf-8", errors="replace", newline="") as stream:
        reader = csv.reader(stream, delimiter=delimiter)
        return "\n".join(" | ".join(cell.strip() for cell in row) for row in reader)


def _read_html(path: Path) -> str:
    soup = BeautifulSoup(_read_text(path), "html.parser")
    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()
    return soup.get_text("\n", strip=True)


def _read_pdf(path: Path) -> tuple[str, dict[str, Any]]:
    reader = PdfReader(str(path))
    pages = [(page.extract_text() or "").strip() for page in reader.pages]
    text = "\n\n".join(page for page in pages if page)
    metadata = {
        "page_count": len(reader.pages),
        "ocr_used": False,
        "warning": None if text else "No embedded text found; OCR is not performed.",
    }
    return text, metadata


def _read_docx(path: Path) -> tuple[str, dict[str, Any]]:
    document = DocxDocument(str(path))
    blocks = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            blocks.append(" | ".join(cell.text.strip() for cell in row.cells))
    return "\n".join(blocks), {
        "paragraph_count": len(document.paragraphs),
        "table_count": len(document.tables),
    }


def parse_file(path: Path, *, max_bytes: int, source: str = "filesystem") -> ParsedDocument:
    if not path.exists() or not path.is_file():
        raise FileNotFoundError(path)
    suffix = path.suffix.casefold()
    if suffix not in SUPPORTED_SUFFIXES:
        raise UnsupportedDocumentError(f"Unsupported file type: {suffix or '<none>'}")
    size = path.stat().st_size
    if size > max_bytes:
        raise ValueError(f"File exceeds max size ({size} > {max_bytes} bytes): {path}")

    metadata: dict[str, Any] = {"size_bytes": size, "suffix": suffix}
    if suffix in {".txt", ".md", ".markdown", ".log"}:
        text = _read_text(path)
    elif suffix in {".json", ".jsonl"}:
        text = _read_json(path)
    elif suffix in {".yaml", ".yml"}:
        text = _read_yaml(path)
    elif suffix in {".csv", ".tsv"}:
        text = _read_delimited(path)
    elif suffix in {".html", ".htm"}:
        text = _read_html(path)
    elif suffix == ".pdf":
        text, extra = _read_pdf(path)
        metadata.update(extra)
    elif suffix == ".docx":
        text, extra = _read_docx(path)
        metadata.update(extra)
    else:  # pragma: no cover - protected by supported suffix check
        raise UnsupportedDocumentError(suffix)

    text = normalise_text(text.replace("\x00", " "))
    if not text:
        raise ValueError(f"Document contains no readable text: {path}")
    content_hash = sha256_text(text)
    source_uri = path.resolve().as_uri()
    return ParsedDocument(
        source=source,
        source_uri=source_uri,
        title=path.name,
        text=text,
        content_hash=content_hash,
        ingest_key=stable_hash(source_uri, content_hash, PARSER_VERSION),
        metadata=metadata,
    )
